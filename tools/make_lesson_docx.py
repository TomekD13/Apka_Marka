# -*- coding: utf-8 -*-
"""Generuje kompletna lekcje (1 temat = 1 plik Word .docx) do korekty.
Tresc w jednokolumnowej tabeli; komorki kolorowane wg poziomu:
  Prosto (base)   -> bez koloru
  Wiecej (extended) -> jasna zielen
  Ekspert (advanced) -> jasny zloty
Naglowki sekcji i Zastosowanie -> ciemny pasek (struktura).
Tekst wersetow wklejany z public/content/pl/bibles/UBG.json.

Uzycie:
  python tools/make_lesson_docx.py 1            # temat o order=1
  python tools/make_lesson_docx.py 1 2 3        # kilka tematow
  python tools/make_lesson_docx.py all          # wszystkie
Wynik: <repo>/lekcje biblijne/NN - <tytul>.docx
"""
import json, os, glob, sys, re
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))   # .../Aplikacja
PL = os.path.join(ROOT, 'public', 'content', 'pl')
OUTDIR = os.path.join(os.path.dirname(ROOT), 'lekcje biblijne')

LEVEL_LABEL = {'base': 'PROSTO', 'extended': 'WIĘCEJ', 'advanced': 'EKSPERT'}
LEVEL_FILL = {'base': None, 'extended': 'E2EFDA', 'advanced': 'FFF2CC'}
LEVEL_TAG = {'base': RGBColor(0x44, 0x54, 0x6A),
             'extended': RGBColor(0x38, 0x76, 0x2C),
             'advanced': RGBColor(0x80, 0x60, 0x00)}
HEADER_FILL = '44546A'        # naglowki sekcji (ciemny slate, bialy tekst)
APP_FILL = 'D9E2F3'           # Zastosowanie (jasny niebieski)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
GREY = RGBColor(0x60, 0x60, 0x60)

CAT = {'basic_practical': 'Studium praktyczne', 'basic_systematic': 'Studium systematyczne',
       'topic': 'Temat biblijny', 'series': 'Seria'}
NOTETYPE = {'syntax': 'składnia', 'verb': 'czasownik', 'background': 'tło', 'textual': 'krytyka tekstu',
            'term': 'termin', 'variant': 'wariant', 'worship': 'uwielbienie', 'other': 'nota'}
LANGNAME = {'grc': 'greka', 'hbo': 'hebrajski', 'he': 'hebrajski', 'gr': 'greka', 'arc': 'aramejski'}


def shade(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill)
    tcPr.append(shd)


def complex_font(run, name):
    run.font.name = name
    rpr = run._element.get_or_add_rPr()
    rf = rpr.find(qn('w:rFonts'))
    if rf is None:
        rf = OxmlElement('w:rFonts'); rpr.append(rf)
    for a in ('w:ascii', 'w:hAnsi', 'w:cs'):
        rf.set(qn(a), name)


def add_cell(table, fill):
    cell = table.add_row().cells[0]
    cell.width = Cm(16)
    if fill:
        shade(cell, fill)
    return cell


def tag_para(cell, level, head):
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(f"[{LEVEL_LABEL[level]}]  ")
    r.bold = True; r.font.color.rgb = LEVEL_TAG[level]; r.font.size = Pt(9)
    r2 = p.add_run(head)
    r2.bold = True; r2.font.size = Pt(11)
    return p


def verse_para(cell, label, text):
    p = cell.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    if label:
        rl = p.add_run(f"{label}  "); rl.bold = True; rl.font.size = Pt(10)
    rt = p.add_run(f"„{text}\""); rt.italic = True; rt.font.size = Pt(10)
    rt.font.color.rgb = RGBColor(0x33, 0x33, 0x33)


def body_para(cell, text):
    p = cell.add_paragraph(); p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text); r.font.size = Pt(11)
    return p


def question_para(cell, text):
    p = cell.add_paragraph(); p.paragraph_format.space_after = Pt(2)
    r = p.add_run("Pytanie: "); r.bold = True; r.italic = True; r.font.size = Pt(10)
    r2 = p.add_run(text); r2.italic = True; r2.font.size = Pt(10)


def original_para(cell, o):
    p = cell.add_paragraph(); p.paragraph_format.space_after = Pt(2)
    ln = LANGNAME.get(o.get('lang', ''), o.get('lang', ''))
    r = p.add_run(f"Forma oryginalna ({ln}): "); r.bold = True; r.font.size = Pt(10)
    r2 = p.add_run(o.get('text', '')); complex_font(r2, 'Times New Roman'); r2.font.size = Pt(12)
    if o.get('translit'):
        r3 = p.add_run(f"  – {o['translit']}"); r3.italic = True; r3.font.size = Pt(10)
        r3.font.color.rgb = GREY


def header_row(table, text, fill=HEADER_FILL, color=WHITE):
    cell = add_cell(table, fill)
    p = cell.paragraphs[0]
    r = p.add_run(text); r.bold = True; r.font.color.rgb = color; r.font.size = Pt(12)


def build(study, series_title, verses):
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Cm(21); sec.page_height = Cm(29.7)
    sec.left_margin = sec.right_margin = Cm(2.5)
    sec.top_margin = sec.bottom_margin = Cm(2.0)
    normal = doc.styles['Normal']; normal.font.name = 'Calibri'; normal.font.size = Pt(11)

    # --- naglowek dokumentu ---
    h = doc.add_heading(f"Temat {study.get('order')}. {study.get('title')}", level=0)
    sub = doc.add_paragraph()
    rs = sub.add_run(f"{series_title}  ·  {CAT.get(study.get('category'), study.get('category',''))}"
                     f"  ·  Czas: Prosto ~{study.get('minutes',{}).get('base','?')} min,"
                     f" Więcej ~{study.get('minutes',{}).get('extended','?')} min")
    rs.italic = True; rs.font.color.rgb = GREY; rs.font.size = Pt(10)
    if study.get('tags'):
        pt = doc.add_paragraph(); rtg = pt.add_run("Tagi: " + ", ".join(study['tags']))
        rtg.italic = True; rtg.font.color.rgb = GREY; rtg.font.size = Pt(10)
    if study.get('summary'):
        ps = doc.add_paragraph(); rsum = ps.add_run("Streszczenie: "); rsum.bold = True
        ps.add_run(study['summary'])

    # --- legenda ---
    doc.add_paragraph().add_run("Legenda poziomów:").bold = True
    leg = doc.add_table(rows=0, cols=1); leg.alignment = WD_TABLE_ALIGNMENT.LEFT
    leg.style = 'Table Grid'
    for lvl, desc in (('base', 'PROSTO – poziom podstawowy (rdzeń, bez koloru)'),
                      ('extended', 'WIĘCEJ – pogłębienie (zielony)'),
                      ('advanced', 'EKSPERT – noty zaawansowane (złoty)')):
        c = add_cell(leg, LEVEL_FILL[lvl]); c.width = Cm(16)
        rr = c.paragraphs[0].add_run(desc); rr.bold = True; rr.font.size = Pt(9)
        rr.font.color.rgb = LEVEL_TAG[lvl]
    doc.add_paragraph()

    # --- glowna tabela lekcji ---
    table = doc.add_table(rows=0, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    table.allow_autofit = False

    for i, s in enumerate(study.get('sections', []), 1):
        header_row(table, f"SEKCJA {i} – {s.get('heading','')}")
        for it in s.get('items', []):
            lvl = it.get('level', 'base')
            cell = add_cell(table, LEVEL_FILL.get(lvl))
            if it.get('type') == 'passage':
                refs = it.get('passage', []) or []
                head = "; ".join(p.get('ref', '') for p in refs)
                tag_para(cell, lvl, head)
                for p in refs:
                    txt = verses.get(p.get('osis', ''))
                    lab = p.get('ref', '') if len(refs) > 1 else ''
                    if txt:
                        verse_para(cell, lab, txt)
                    else:
                        verse_para(cell, lab, f"[brak tekstu w module UBG dla {p.get('osis','')}]")
                if it.get('comment'):
                    body_para(cell, it['comment'])
                for q in it.get('questions', []) or []:
                    question_para(cell, q.get('text', ''))
            else:  # note
                nt = NOTETYPE.get(it.get('noteType'), it.get('noteType', ''))
                tag_para(cell, lvl, f"Nota ({nt}): {it.get('label','')}")
                if it.get('content'):
                    body_para(cell, it['content'])
                for o in it.get('original', []) or []:
                    original_para(cell, o)
                for q in it.get('questions', []) or []:
                    question_para(cell, q.get('text', ''))

    # --- zastosowanie ---
    app = study.get('application', {}) or {}
    if app.get('text') or app.get('challenge'):
        header_row(table, "ZASTOSOWANIE")
        cell = add_cell(table, APP_FILL)
        if app.get('text'):
            body_para(cell, app['text'])
        if app.get('challenge'):
            p = cell.add_paragraph(); r = p.add_run("Wyzwanie na tydzień: "); r.bold = True
            p.add_run(app['challenge'])

    # --- stopka ---
    doc.add_paragraph()
    f = doc.add_paragraph()
    rf = f.add_run("Tekst Pisma: Uwspółcześniona Biblia Gdańska (UBG). "
                   "Materiał roboczy do korekty – formy oryginalne (gr./hebr.) i tło historyczne do weryfikacji.")
    rf.italic = True; rf.font.size = Pt(9); rf.font.color.rgb = GREY
    return doc


def safe_name(s):
    s = re.sub(r'[\\/:*?"<>|]', '', s).strip().rstrip('.')
    return re.sub(r'\s+', ' ', s)


def main():
    args = sys.argv[1:] or ['1']
    idx = json.load(open(os.path.join(PL, 'index.json'), encoding='utf-8'))
    series_title = {s['id']: s['title'] for s in idx.get('series', [])}
    ubg = json.load(open(os.path.join(PL, 'bibles', 'UBG.json'), encoding='utf-8'))
    verses = ubg.get('verses', {})

    studies = {}
    for fp in glob.glob(os.path.join(PL, 'studies', '*.json')):
        d = json.load(open(fp, encoding='utf-8'))
        studies[d.get('order')] = d

    if args == ['all']:
        orders = sorted(studies)
    else:
        orders = [int(a) for a in args]

    os.makedirs(OUTDIR, exist_ok=True)
    for o in orders:
        d = studies.get(o)
        if not d:
            print(f"  pomijam order={o} (brak)"); continue
        doc = build(d, series_title.get(d.get('seriesId'), ''), verses)
        fn = f"{o:02d} - {safe_name(d.get('title',''))}.docx"
        out = os.path.join(OUTDIR, fn)
        doc.save(out)
        print(f"OK  {fn}")
    print(f"\nKatalog: {OUTDIR}")


if __name__ == '__main__':
    main()
