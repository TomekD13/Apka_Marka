# -*- coding: utf-8 -*-
"""Generuje dokument Word z pozycjami czekajacymi na decyzje autora.

Zrodlo: sekcja „Do decyzji" w KOREKTA-PL-do-rozstrzygniecia.md + pelny tekst pola
ze studium + tekst wersetu z BE.json, zeby dalo sie decydowac bez zagladania w kod.
Kazda pozycja konczy sie pustym polem „TWOJA DECYZJA" do wpisania odpowiedzi.

Uzycie:
  python tools/make_decyzje_docx.py
Wynik: <repo>/KOREKTA-PL-do-rozstrzygniecia.docx
"""
import io, json, os

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))          # .../Aplikacja
REPO = os.path.dirname(ROOT)
STUDIES = os.path.join(ROOT, 'public', 'content', 'pl', 'studies')
BIBLE = os.path.join(ROOT, 'public', 'content', 'pl', 'bibles', 'BE.json')
RAPORT = os.path.join(REPO, 'KOREKTA-PL-do-rozstrzygniecia.md')
OUT = os.path.join(REPO, 'KOREKTA-PL-do-rozstrzygniecia.docx')

HEAD_FILL = '44546A'        # pasek naglowka pozycji
JEST_FILL = 'F2F2F2'        # tekst, ktory dzis stoi w aplikacji
WERSET_FILL = 'EAF1F8'      # werset z BE
REKO_FILL = 'E2EFDA'        # moja rekomendacja
DECYZJA_FILL = 'FFF2CC'     # puste pole na odpowiedz
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
GREY = RGBColor(0x59, 0x59, 0x59)

# ---------------------------------------------------------------- rekomendacje
# klucz: (plik, gdzie, poczatek pola „Jest")  ->  (werdykt, uzasadnienie, nowe brzmienie)
R = {}


def reko(plik, gdzie, jest, werdykt, uzas, nowe=''):
    R[(plik, gdzie, jest.strip())] = (werdykt, uzas, nowe)


def znajdz_reko(plik, gdzie, jest):
    """Dopasowanie po (plik, gdzie) + poczatku pola „Jest".

    Prefiks w tabeli wyzej bywa krotszy niz tekst w raporcie, wiec porownujemy
    tyle znakow, ile maja obie strony. Gdy do jednego miejsca jest kilka zgloszen
    (jedno zdanie, dwa zarzuty), wygrywa najdluzsze wspolne dopasowanie."""
    jest = (jest or '').strip()
    kandydaci = [(k, v) for k, v in R.items() if k[0] == plik and k[1] == gdzie]
    if not kandydaci:
        return '', '', ''
    if len(kandydaci) == 1:
        return kandydaci[0][1]
    najlepszy, ile = None, -1
    for k, v in kandydaci:
        n = min(len(k[2]), len(jest))
        wspolne = 0
        while wspolne < n and k[2][wspolne] == jest[wspolne]:
            wspolne += 1
        if wspolne > ile:
            najlepszy, ile = v, wspolne
    return najlepszy


reko('bestie-i-baranek.json', 's2/u5/comment', '„oczami jak u człowieka i ustami mówiącymi',
     'PRZEREDAGOWAĆ',
     'Korektor ma rację co do cytatu, ale jego wersja nie wchodzi po przyimku „z". '
     'Wystarczy rozbić zdanie na dwa człony.',
     'Spomiędzy rogów wyrasta „mały róg", który ma „oczy podobne do ludzkich oczu i usta, '
     'które mówiły zuchwałe rzeczy". To nie kolejne imperium, lecz moc o charakterze '
     'religijno-politycznym, działająca słowem.')

reko('bestie-i-baranek.json', 's2/u6/comment', '„myśli o zmianie czasów i Prawa"',
     'PRZEREDAGOWAĆ',
     'BE ma tu czas przyszły („będzie wypowiadał", „będzie chciał zmienić"), więc całe '
     'wyliczenie trzeba przestawić na przyszły – wtedy oba cytaty wchodzą dosłownie. '
     'Ta sama poprawka załatwia pozycję niżej („czas, dwa czasy i pół czasu").',
     'Mały róg „będzie wypowiadał słowa przeciw Najwyższemu" i „będzie również chciał zmienić '
     'czasy i prawo", a święci zostaną wydani w jego ręce „aż do czasu, dwóch czasów i pół czasu". '
     'To krytyka systemu i fałszywego nauczania - nie potępienie ludzi.')

reko('bestie-i-baranek.json', 's2/u6/comment', '„czas, czasy i pół czasu"',
     'PRZEREDAGOWAĆ (razem z pozycją wyżej)',
     'BE ma „dwóch czasów", nie „czasy". Wchodzi jednym ruchem razem z poprawką powyżej – '
     'to jedno i to samo zdanie.',
     '(patrz brzmienie zaproponowane w pozycji poprzedniej)')

reko('biblia.json', 's3/u10/comment', '„Żadne proroctwo Pisma nie podlega dowolnemu',
     'TAK',
     'Zdanie zaczyna się cytatem, więc wielka litera jest poprawna typograficznie '
     'i nie narusza wierności wobec BE.',
     '„Wszelkie proroctwo Pisma nie może być dowolnie wykładane". Pismo wyjaśnia samo siebie - '
     'jasne fragmenty pomagają zrozumieć trudniejsze, a żaden werset nie znaczy tego, '
     'co nam wygodnie.')

reko('co-po-smierci.json', 's1/u4/comment', '„To nie umarli chwalą Pana… lecz my będziemy',
     'DO WYBORU (A lub B)',
     'Kontrast „to nie umarli… lecz my" jest sercem komentarza, a jego druga połowa to Ps 115,18, '
     'którego nie ma w odnośniku. Sprawdziłem w źródle BE: w. 18 brzmi „ale to my błogosławimy '
     'PANA, teraz i na wieki". Propozycja korektora po prostu wycina kontrast – to zmiana sensu, '
     'nie korekta.',
     'A (rekomendowane): rozszerzyć odnośnik do Ps 115,17-18 i zacytować dosłownie – '
     '„To nie umarli wychwalają PANA (...), ale to my błogosławimy PANA, teraz i na wieki". '
     'Wymaga przebudowy BE.json przez build_bible_be.py.\n'
     'B: skrócić cytat do w. 17 i kontrast oddać własnymi słowami, bez cudzysłowu.')

reko('czy-ktos-sie-o-mnie-troszczy.json', 's2/n3/q18', 'miłości wybierającej (a nie odwzajemnianej)',
     'TAK',
     'Zarzut trafny – „wybierająca" nie jest przeciwieństwem „odwzajemnianej". '
     'Propozycja oddaje to, o co w agapē faktycznie chodzi.',
     'Jak rozumienie *agapē* jako miłości wybierającej dobro drugiego '
     '(a nie uzależnionej od wzajemności) zmienia obraz Boga?')

reko('czym-jest-grzech.json', 's4/u14/comment', '„umrzeć dla grzechu, a żyć dla sprawiedliwości"',
     'PRZEREDAGOWAĆ',
     'Cytat trzeba wziąć z BE w całości, razem ze spójnikiem „abyśmy" – inaczej zdanie '
     'się nie klei.',
     'Jezus „w swoim ciele poniósł nasze grzechy na drzewo", „abyśmy martwi dla grzechu, '
     'żyli dla sprawiedliwości". Krzyż nie tylko przebacza, ale uwalnia.')

reko('czym-jest-grzech.json', 's5/u17/comment', '„nie poczytuje ludziom ich upadków"',
     'PRZEREDAGOWAĆ',
     'BE ma imiesłów „nie poczytując", więc trzeba dołożyć zdanie główne z w. 19.',
     'Sercem krzyża jest „pojednanie" - w Chrystusie Bóg pojednał świat ze sobą, '
     '„nie poczytując ludziom ich grzechów", i przywraca zerwaną więź. '
     'Cel to relacja, nie tylko czyste konto.')

reko('duch-swiety.json', 's1/u5/comment', '„przyczynia się za nami"',
     'PRZEREDAGOWAĆ',
     '„przyczynia się za nami" to brzmienie UBG/BW – w BE nie pada. Propozycja korektora '
     'powtarza tę samą frazę dwa razy w jednym zdaniu; lepiej wziąć drugi czasownik z BE '
     '(„wspiera nas w naszej niemocy").',
     'Duch „wspiera nas w naszej niemocy" i „wstawia się za nami w błaganiach, których nie można '
     'wyrazić słowami". To portret kogoś bliskiego, kto rozumie i oręduje, nie bezosobowej siły.')

reko('duch-swiety.json', 's3/n6/label', '„Abba": najczulsze słowo „Tato"',
     'TAK (z drobną zmianą)',
     'Zarzut zasadny: utożsamienie „Abba" z dziecięcym „Tato" to znany błąd popularyzatorski '
     '(podważył go James Barr, „Abba isn\'t Daddy"). „Abba" jest zwrotem poufałym i rodzinnym, '
     'ale nie dziecinnym – używali go także dorośli synowie. Proponuję zachować ciepło obrazu, '
     'bez spornego „Tato".',
     '„Abba": poufały, rodzinny zwrot „Ojcze" (Rz 8,15)')

reko('duch-swiety.json', 's4/u15/comment', '„Otrzymacie moc Ducha Świętego... i będziecie',
     'PRZEREDAGOWAĆ',
     'Cytat wziąć z BE w całości; kropka zostaje na końcu cytatu, więc zdania się nie zlewają.',
     '„Gdy Duch Święty zstąpi na was, weźmiecie Jego moc i będziecie Moimi świadkami". '
     'Duch daje moc nie do popisu, lecz do życia i mówienia o Jezusie.')

reko('gdy-upadam.json', 's1/u2/comment', '„doświadczony był we wszystkim… podobnie jak my',
     'PRZEREDAGOWAĆ',
     'BE ma dopełniacz („doświadczonego"), bo zdanie zależy od „nie mamy arcykapłana". '
     'Trzeba więc zacytować od tego miejsca, a nie wstawiać cytat po „Jezus".',
     'Nie mamy arcykapłana, „który nie mógłby współczuć w naszych słabościach", lecz „podobnie '
     'jak my doświadczonego we wszystkim, oprócz grzechu". Jezus rozumie walkę z pokusą '
     'od wewnątrz - nie patrzy na nią z dystansu.')

reko('gdy-upadam.json', 's2/u4/comment', '„Bóg daje nam zwycięstwo przez Pana naszego',
     'PRZEREDAGOWAĆ',
     'W BE podmiotem zdania jest dziękczynienie, więc „Bóg" trzeba wyprowadzić przed cytat, '
     'a nie wciągać do środka.',
     'Zwycięstwo nad grzechem nie jest naszym osiągnięciem - to Bóg „daje nam zwycięstwo przez '
     'naszego Pana Jezusa Chrystusa". Walczymy z Jego mocy, nie własnej siły woli.')

reko('gdy-upadam.json', 's3/n6/content', 'Północ i południe mają punkty graniczne',
     'NIE – zarzut chybiony',
     'Korektor mierzy odległość, a nota mówi o kierunku. Asymetria jest realna: idąc na północ '
     'mijasz biegun i od razu idziesz na południe; idąc na wschód, nigdy nie zaczniesz iść '
     'na zachód. Zamiana zdania niepotrzebnie osłabiłaby notę. Jeśli chcesz dmuchać na zimne, '
     'wystarczy dopisać pół zdania, że to obraz poetycki, a nie pomiar.',
     'Bez zmian (ewentualnie dopisek: „To obraz poetycki, nie pomiar odległości.")')

reko('godzina-sadu.json', 's2/u8/comment', 'Sędzia ma twarz naszego Obrońcy',
     'TAK',
     'Zarzut trafny – i, co ważniejsze, to samo studium rozstrzyga sprawę we własnych notach: '
     'n7 mówi „Ojciec jest Sędzią, a Chrystus stoi po stronie sądzonego", n8 „Chrystus «nie sądzi» '
     'w tej fazie". Obecne zdanie przeczy więc notom obok. Przy okazji warto zrównać cytat '
     'z 1 J 2,1 z BE, która ma „Orędownika przed Ojcem", nie „Rzecznika u Ojca".',
     'W tej sali nie jesteśmy sami: „mamy Orędownika przed Ojcem – sprawiedliwego Jezusa '
     'Chrystusa". Przed Sędzią stoi nasz Obrońca, który oddał za nas życie.')

reko('godzina-sadu.json', 's2/u8/q8', 'ten, kto cię sądzi, jest zarazem tym, kto za',
     'TAK',
     'To samo co wyżej – pytanie musi trzymać się podziału ról, który studium samo wprowadza.',
     'Co zmienia świadomość, że ten, kto za ciebie umarł, jest zarazem twoim Obrońcą?')

reko('historia-swiata-w-proroctwie.json', 'application/text', 'na wieki naprzód nazwano bieg mocarstw',
     'TAK',
     'Zarzut trafny: Dn 2 nazywa z imienia tylko Babilon, a ten już panował. Precyzja tu wzmacnia '
     'argument, bo łatwiej go obronić w rozmowie z kimś, kto zna tekst.',
     '…Daniel 2 to jeden z najmocniejszych argumentów za zaufaniem do Biblii: z wyprzedzeniem '
     'zapowiedziano bieg mocarstw, a historia to potwierdziła. (dalej bez zmian)')

reko('historia-swiata-w-proroctwie.json', 's2/u6/q6', 'nazwał kolejne mocarstwa na wieki przed ich',
     'TAK',
     'To samo zastrzeżenie co wyżej, w pytaniu.',
     'Co mówi o Bogu fakt, że zapowiedział następstwo kolejnych mocarstw przed ich panowaniem?')

reko('jak-rozpoznawac-boza-wole.json', 's4/u14/comment', 'wewnętrzne potwierdzenie lub niepokój wobec',
     'TAK',
     'Flp 4,7 mówi o pokoju, który strzeże serca – nie o pokoju jako detektorze woli Bożej. '
     'To popularna, ale nieuprawniona ekstrapolacja i w praktyce prowadzi do rozeznawania '
     'po samopoczuciu, wbrew sola scriptura.',
     '„We wszystkim... niech wasze prośby będą znane Bogu, a pokój Boży... ustrzeże wasze serca." '
     'Modlitwa o decyzję często przynosi nie tyle słyszalną odpowiedź, ile „pokój Boży, który '
     'przewyższa wszelkie zrozumienie" - ten, który „strzeże serc i myśli". To nie wskaźnik '
     'kierunku, lecz oparcie w czasie decydowania.')

reko('kim-jest-jezus.json', 's2/u4/comment', 'dlatego słuchacze chcą Go ukamienować za',
     'TAK (z uzupełnieniem)',
     'J 8,58 nie podaje motywu – podaje go dopiero J 10,33 („za bluźnierstwo"). Najprościej '
     'opisać reakcję, a motyw wesprzeć odsyłaczem, zamiast wkładać go w werset.',
     'Jezus nie mówi „byłem, zanim był Abraham", lecz „JESTEM". Świadomie sięga po imię Boga '
     'z Wj 3,14 - a reakcją słuchaczy jest próba ukamienowania Go (J 8,59; por. J 10,33).')

reko('kim-jest-jezus.json', 's3/n9/content', 'Czasownik „oddawać pokłon, cześć" (gr. *pros',
     'TAK (przeredagować całą notę)',
     'Zarzut jest poważny: proskyneō bywa użyte wobec ludzi (Mt 18,26; Ap 3,9), więc twierdzenie '
     'kategoryczne jest łatwe do obalenia i przewraca całą notę. Argument da się jednak uratować – '
     'nie na słowniku, lecz na reakcji aniołów i apostołów, którzy czci odmawiają, a Jezus nie.',
     'Czasownik „oddawać pokłon, cześć" (gr. *proskyneō*) sam w sobie może oznaczać zwykły hołd, '
     'ale w kontekście religijnym opisuje cześć należną wyłącznie Bogu. Dlatego aniołowie '
     '(Ap 22,8-9) i apostołowie (Dz 10,25-26) stanowczo odmawiają jej przyjęcia. Jezus natomiast '
     'cześć przyjmuje i nigdy jej nie odrzuca (Mt 14,33; 28,9.17; J 9,38). Przy konsekwentnym '
     'monoteizmie Biblii jest to ciche, lecz mocne świadectwo, że Jezus jest Bogiem.')

reko('lek-stres-pokoj.json', 's2/n4/content', '(Łk 19,35)',
     'NIE – bezprzedmiotowe',
     'Propozycja jest identyczna z tekstem, który już tam stoi. Sprawdziłem sam odnośnik: '
     'Łk 19,35 rzeczywiście ma formę czasownika epiriptō (ἐπιρίψαντες) o zarzuceniu szat na oślę. '
     'Nota jest poprawna.',
     'Bez zmian.')

reko('lek-stres-pokoj.json', 's2/u6/comment', '„Zrzuć swój ciężar na Pana, a On cię podtrzyma."',
     'PRZEREDAGOWAĆ (ale nie wg korektora)',
     'Propozycja korektora cytuje zupełnie inny werset – patrzył jeszcze na stary, błędny '
     'odnośnik. Odnośnik jest już poprawiony i obok stoi BE Ps 55,23; wystarczy zrównać z nim '
     'cytat. Obraz plecaka w dalszej części zostaje, bo to już nie cytat.',
     '„Zrzuć swoją troskę na PANA, a On cię podtrzyma". Obraz jest fizyczny: nie da się '
     'jednocześnie nieść plecaka i go oddać. Trzymanie i ufanie nawzajem się wykluczają.')

reko('lek-stres-pokoj.json', 's2/u8/comment', '„Bóg jest dla nas ucieczką i siłą... dlatego',
     'PRZEREDAGOWAĆ (ale nie wg korektora)',
     'Jak wyżej – korektor cytował spod starego odnośnika. Po poprawce obok stoi BE Ps 46,2-3 '
     'i cytat wystarczy z nim zrównać.',
     '„Bóg jest naszą ucieczką i mocą (...). Dlatego się nie zlękniemy, choćby zatrzęsła się '
     'ziemia". Pokój nie zależy od tego, czy świat jest stabilny, lecz od tego, kto jest '
     'naszą skałą.')

reko('nawrocenie-i-nowe-zycie.json', 's1/u1/comment', '„nawróćcie się i wierzcie w Ewangelię"',
     'ZALEŻY OD DECYZJI O Mk 1,15 (sekcja A)',
     'Sprawdziłem tekst BE w źródle: „Nadszedł czas, Królestwo Boga jest już blisko, nawróćcie '
     'się i wierzcie w Ewangelię". Jeśli zdejmiemy fallback z Biblii Warszawskiej, obecny cytat '
     'jest dosłownie zgodny z BE i NIE TRZEBA ZMIENIAĆ NIC. Jeśli fallback zostaje – trzeba '
     'wpisać „upamiętajcie się i wierzcie ewangelii".',
     'Rekomendacja: wrócić do BE (patrz sekcja A) – wtedy ta pozycja odpada sama.')

reko('prorocze-wyliczenia.json', 's1/n4/content', 'skoro liczba sięga „czasu końca" (Dn 8,17.19)',
     'TAK',
     '„nie mogą… muszą" to logiczny nadmiar – z tego, że wizja sięga czasu końca, nie wynika '
     'koniecznie przelicznik. Nazwanie tego odczytem historycystycznym jest uczciwsze '
     'i odporniejsze na zarzut, a wniosku nie osłabia.')

reko('prorocze-wyliczenia.json', 's2/n6/content', 'skoro tuż obok mowa o „siedemdziesięciu latach"',
     'TAK',
     'Zarzut trafny: Dn 9,24 mówi o siedemdziesięciu TYGODNIACH, nie o samej liczbie 70. '
     'Sąsiedztwo tekstów jest przesłanką wspierającą, nie dowodem – i tak trzeba to nazwać.')

reko('prorocze-wyliczenia.json', 's3/n8/content', 'a skoro są odcięte z początku, obie linie',
     'TAK',
     'Czasownik chatak mówi „odciąć", ale nie mówi „z początku". To założenie interpretacyjne '
     'i lepiej je nazwać po imieniu – inaczej w oczach czytelnika, który to sprawdzi, '
     'wygląda na przemyt.')

reko('prorocze-wyliczenia.json', 's3/u11/comment', 'Lata później anioł wraca',
     'NIE – zarzut błędny',
     'Dn 8 to trzeci rok Belszazara, Dn 9 pierwszy rok Dariusza – dzieli je kilkanaście lat. '
     '„Lata później" jest poprawne.',
     'Bez zmian.')

reko('samotnosc-smutek-depresja.json', 's2/u5/comment', '„Bliski jest Pan dla skruszonych w sercu',
     'PRZEREDAGOWAĆ (ale nie wg korektora)',
     'Korektor cytował spod starego odnośnika. Po poprawce obok stoi BE Ps 34,19 i cytat '
     'wystarczy z nim zrównać. Etykietę noty w tym studium poprawiłem już na „złamani sercem" – '
     'teraz jedno z drugim będzie się zgadzać.',
     '„PAN jest blisko tych, których serce jest złamane i wybawia strapionych na duchu". '
     'Bóg nie czeka, aż się pozbieramy - najbliżej jest właśnie wtedy, gdy jesteśmy w kawałkach.')

reko('samotnosc-smutek-depresja.json', 's2/u9/comment', '„Policz moje tułaczki, zbierz moje łzy',
     'PRZEREDAGOWAĆ (ale nie wg korektora)',
     'Jak wyżej – propozycja korektora dotyczy innego wersetu. BE Ps 56,9 ma czas przeszły '
     '(„Policzyłeś… zebrałeś"), co zresztą brzmi mocniej: to już się stało.',
     '„Policzyłeś dni mojej tułaczki, moje łzy zebrałeś do swego bukłaka. Czy nie są spisane '
     'w Twojej księdze?" Żadna twoja łza nie jest niezauważona ani zmarnowana.')

reko('skad-zlo-i-cierpienie.json', 's3/u8/comment', 'pięciokrotne „wstąpię... będę równy',
     'NIE – zarzut błędny (cytat już poprawiony)',
     'W BE Iz 14,13-14 jest PIĘĆ deklaracji: „wstąpię", „wyniosę", „zamieszkam", „wstąpię", '
     '„stanę się". „Pięciokrotne" zostaje; poprawiono tylko samo brzmienie cytatu.',
     'Bez zmian.')

reko('swiatynia-i-przymierze.json', 's1/u3/comment', '„Twoja droga, Boże, jest w świątyni"',
     'DO WYBORU (A lub B) – realny problem',
     'BE oddaje Ps 77,14 jako „Boże, święta jest Twoja droga" – nie „w świątyni". Tak samo BW, BT '
     'i UBG. Odczyt „w świątyni" (hebr. בַּקֹּדֶשׁ ba-kodesz może znaczyć „w świętości" albo '
     '„w miejscu świętym") idzie za KJV i jest w polszczyźnie mniejszościowy. Cały komentarz stoi '
     'na tym jednym odczycie, więc samą podmianą cytatu się tego nie załatwi.',
     'A (rekomendowane): przeredagować tak, by myśl nie wisiała na spornym tłumaczeniu – '
     '„Przybytek przedstawia drogę, jaką Bóg zbawia człowieka. Hebrajskie «Twoja droga jest '
     "ba-kodesz» bywa oddawane «święta» albo «w miejscu świętym». Każdy element przybytku "
     'coś znaczy."\n'
     'B: oprzeć jednostkę na innym wersecie, np. Wj 25,8-9 – „Zbudują Mi też świątynię, '
     'abym zamieszkał wśród nich".')

reko('szabat.json', 's2/n6/q20', '„znanie" Boga jako bliska relacja',
     'NIE – bezprzedmiotowe',
     'Propozycja jest identyczna z tekstem. Składnia jest poprawna: podmiotem „różni się" jest '
     '„znanie Boga jako bliska relacja", więc mianownik stoi na miejscu.',
     'Bez zmian.')

reko('wdziecznosc-radosc-sens.json', 's2/u5/comment', '„Zawsze się radujcie, nieustannie się módl',
     'PRZEREDAGOWAĆ (razem z pozycją niżej)',
     'To jedyne miejsce, w którym zmiana przekładu uderza w samą myśl komentarza: BE ma '
     '„za wszystko dziękujcie", a cała uwaga stoi na rozróżnieniu „w każdym położeniu" kontra '
     '„za każde położenie". Rozróżnienie jest egzegetycznie słuszne – grecki zwrot en panti '
     'znaczy „we wszystkim", nie „za wszystko" – więc nie trzeba go porzucać, tylko oprzeć '
     'na grece zamiast na polskim przekładzie.',
     '„Zawsze się radujcie, nieustannie się módlcie, za wszystko dziękujcie". Uwaga na słowa: '
     'grecki zwrot *en panti* znaczy dosłownie „we wszystkim", czyli w każdym położeniu - '
     'a nie „za każde położenie". To wdzięczność w trudzie, nie za zło.')

reko('wdziecznosc-radosc-sens.json', 's2/u5/comment', 'Uwaga na słowa: dziękować mamy „w każdym',
     'PRZEREDAGOWAĆ (razem z pozycją wyżej)',
     'To drugie zgłoszenie do tego samego zdania – załatwia je brzmienie zaproponowane powyżej. '
     'Propozycja korektora („w podanym przekładzie czytamy…") gubi całą myśl uwagi.',
     '(patrz brzmienie zaproponowane w pozycji poprzedniej)')

reko('zbawienie-za-darmo.json', 's3/u13/comment', 'nawet „ja sam" nie jestem w niej wyjątkiem',
     'TAK',
     'Zarzut trafny i ważny teologicznie: obecne zdanie ociera się o zaprzeczenie wolnej woli, '
     'a to idzie wprost przeciw osi wielkiego boju. Rz 8,38-39 wylicza moce zewnętrzne – '
     'nie własną decyzję człowieka.',
     '„Ani śmierć, ani życie (...), ani żadne inne stworzenie nie zdoła nas odłączyć od miłości '
     'Boga, która jest w Chrystusie Jezusie". Lista jest wyczerpująca - żadne stworzenie ani '
     'zewnętrzna moc nie są w stanie wyrwać nas z tej miłości.')

reko('zycie-w-bozej-rodzinie.json', 's4/u14/comment', 'Przed Wieczerzą Jezus umywa uczniom nogi',
     'PRZEREDAGOWAĆ',
     'J 13,2-4 mówi „Podczas wieczerzy… wstał od stołu". Obecne zdanie miesza scenę biblijną '
     'z porządkiem liturgicznym ADS, w którym obrzęd pokory faktycznie poprzedza Wieczerzę '
     'Pańską. Warto te dwie rzeczy rozdzielić, żeby nic nie zginęło.',
     'Podczas Wieczerzy Jezus wstaje od stołu i umywa uczniom nogi - czynność niewolnika. '
     'W praktyce Kościoła obrzęd pokory poprzedza Wieczerzę Pańską i nadaje jej ton: '
     'spotkanie z Bogiem, który służy.')

# --------------------------------------------------------- decyzje zasadnicze
ZASADNICZE = [
    ('A1. Mk 1,15 – fallback z Biblii Warszawskiej',
     'Werset idzie dziś z BW („Wypełnił się czas… upamiętajcie się i wierzcie ewangelii"), bo BE '
     'ma „Nadszedł czas", co zaciera dokonaność ważną w studium „prorocze-wyliczenia". Skutek '
     'uboczny: w studium o nawróceniu czytelnik dostaje archaiczne „upamiętajcie się".',
     'Pełny tekst BE (sprawdzony w źródle): „Mówił: Nadszedł czas, Królestwo Boga jest już blisko, '
     'nawróćcie się i wierzcie w Ewangelię."',
     'Wrócić do BE. Argument o wyznaczonym czasie oprzeć na Ga 4,4, gdzie BE ma „Kiedy jednak '
     'nadeszła pełnia czasu, Bóg posłał swego Syna". Zysk potrójny: znika archaizm, pozycja 25 '
     'z listy poniżej odpada sama (cytat w „nawrocenie-i-nowe-zycie" jest wtedy dosłownie zgodny '
     'z BE), a w aplikacji zostają 4 fallbacki zamiast 5.\n'
     'Robota: usunąć Mark.1.15 z FALLBACKS w build_bible_be.py, przebudować BE.json, '
     'przeredagować argument w „prorocze-wyliczenia".'),

    ('A2. Forma hebrajska przy Dn 9,24',
     'Nota o „siedemdziesięciu siódemkach" w „prorocze-wyliczenia" ma w polu original zapis '
     'שִׁבְעִים שִׁבְעִים / šiḇʿîm šiḇʿîm, czyli dwa razy ten sam liczebnik („siedemdziesiąt '
     'siedemdziesiąt"), podczas gdy komentarz mówi „siedemdziesiąt siódemek". Pierwszy wyraz '
     'powinien być rzeczownikiem „tygodnie/siódemki".',
     '',
     'Nie zgadywać. Do czasu sprawdzenia w wydaniu krytycznym (BHS) usunąć pole original '
     'i zostawić samo polskie omówienie, a w meta ustawić needsReview = true. '
     'CLAUDE.md mówi wprost: formy oryginalne tylko pewne, niepewne – pomijać.'),

    ('A3. Prawa do Biblii Ekumenicznej',
     'W aplikacji jest 779 fragmentów przekładu chronionego prawem autorskim '
     '(© Towarzystwo Biblijne w Polsce). Nota licencyjna jest widoczna w „O aplikacji".',
     '',
     'Przed publikacją na www rozważyć wystąpienie do Towarzystwa Biblijnego w Polsce o zgodę. '
     'To decyzja poza moim zasięgiem – potrzebna twoja.'),

    ('A4. Numeracja Psalmów w pozostałych językach',
     'Ten sam błąd numeracji Psalmów, który poprawiono w PL (Ps 55, 46, 34, 56, 77 – cytat '
     'wskazywał werset obok), siedzi nadal w EN/DE/ES/FR/PT/SW/UK. Zgodnie z twoją decyzją '
     'nie ruszałem innych języków.',
     'UWAGA: nie wolno użyć *_refs_from_pl.py – te skrypty wzięłyby numerację BE (hebrajską) '
     'i przeniosły ją do przekładów w numeracji KJV.',
     'Poprawić przy najbliższym przeglądzie danego języka, ręcznie albo osobnym skryptem '
     'operującym na osis, nie na polskim ref. Nie jest to pilne – błąd dotyczy pięciu miejsc.'),
]


# ------------------------------------------------------------------- narzedzia
def shade(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill)
    tcPr.append(shd)


def komorka(table, fill=None):
    cell = table.add_row().cells[0]
    cell.width = Cm(17)
    if fill:
        shade(cell, fill)
    return cell


def wiersz(cell, etykieta, tekst, kursywa=False, rozmiar=10, kolor=None, pierwszy=True):
    p = cell.paragraphs[0] if pierwszy and not cell.paragraphs[0].runs else cell.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    if etykieta:
        r = p.add_run(etykieta + '  ')
        r.bold = True
        r.font.size = Pt(9)
        r.font.color.rgb = GREY
    for i, linia in enumerate(tekst.split('\n')):
        if i:
            p = cell.add_paragraph()
            p.paragraph_format.space_after = Pt(3)
        r = p.add_run(linia)
        r.font.size = Pt(rozmiar)
        r.italic = kursywa
        if kolor:
            r.font.color.rgb = kolor
    return p


def naglowek_pozycji(table, numer, tytul):
    cell = komorka(table, HEAD_FILL)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run('%s   %s' % (numer, tytul))
    r.bold = True
    r.font.size = Pt(11)
    r.font.color.rgb = WHITE


def pole_decyzji(table, ile_linii=4):
    cell = komorka(table, DECYZJA_FILL)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run('TWOJA DECYZJA')
    r.bold = True
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0x80, 0x60, 0x00)
    for _ in range(ile_linii):
        q = cell.add_paragraph()
        q.paragraph_format.space_after = Pt(2)
        q.add_run('').font.size = Pt(11)


def nowa_tabela(doc):
    t = doc.add_table(rows=0, cols=1)
    t.style = 'Table Grid'
    return t


# --------------------------------------------------------------------- wejscie
def wczytaj_pozycje():
    lines = io.open(RAPORT, encoding='utf-8').read().split('\n')
    start = next(i for i, l in enumerate(lines) if l.startswith('## Do decyzji'))
    end = next(i for i, l in enumerate(lines) if i > start and l.startswith('## Zgłoszone'))
    poz, cur = [], None
    for l in lines[start:end]:
        if l.startswith('### '):
            if cur:
                poz.append(cur)
            plik, _, gdzie = l[4:].partition(' – ')
            cur = {'plik': plik.strip(), 'gdzie': gdzie.strip().strip('`')}
        elif cur is not None:
            for et, k in (('- **Jest:** ', 'jest'), ('- **Problem:** ', 'problem'),
                          ('- **Propozycja:** ', 'propozycja')):
                if l.startswith(et):
                    cur[k] = l[len(et):].strip()
    if cur:
        poz.append(cur)
    return poz


def kontekst(poz, verses):
    """Dokleja pelny tekst pola, tytul studium i wersety BE."""
    study = json.load(io.open(os.path.join(STUDIES, poz['plik']), encoding='utf-8'))
    poz['tytul'] = study.get('title', '')
    czesci = poz['gdzie'].split('/')
    poz['tekst'], poz['refs'] = '(nie znaleziono)', []
    if czesci[0] == 'application':
        poz['tekst'] = (study.get('application') or {}).get(czesci[1], '')
        return poz
    for sec in study.get('sections', []):
        if sec.get('id') != czesci[0]:
            continue
        for it in sec.get('items', []):
            if it.get('id') != czesci[1]:
                continue
            key = czesci[-1]
            if key.startswith('q') and key[1:].isdigit():
                poz['tekst'] = next((q.get('text', '') for q in it.get('questions', []) or []
                                     if q.get('id') == key), '(nie znaleziono pytania)')
            else:
                poz['tekst'] = it.get(key, '(brak pola)')
            poz['refs'] = [(p.get('ref'), verses.get(p.get('osis', ''), '(brak w BE.json)'))
                           for p in (it.get('passage') or [])]
    return poz


# ----------------------------------------------------------------------- glowna
def main():
    verses = json.load(io.open(BIBLE, encoding='utf-8'))['verses']
    pozycje = [kontekst(p, verses) for p in wczytaj_pozycje()]

    doc = Document()
    st = doc.styles['Normal']
    st.font.name = 'Calibri'
    st.font.size = Pt(11)
    for s in doc.sections:
        s.left_margin = s.right_margin = Cm(2)

    h = doc.add_paragraph()
    r = h.add_run('Korekta wersji PL – do rozstrzygnięcia')
    r.bold = True
    r.font.size = Pt(20)
    p = doc.add_paragraph()
    r = p.add_run('Biblia Ekumeniczna · stan na 2026-08-21 · %d pozycji + 4 decyzje zasadnicze'
                  % len(pozycje))
    r.font.size = Pt(10)
    r.font.color.rgb = GREY

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    r = p.add_run('Jak z tego korzystać. ')
    r.bold = True
    r.font.size = Pt(10)
    r = p.add_run('Każda pozycja pokazuje: co dziś stoi w aplikacji, jaki werset czytelnik widzi '
                  'obok, co zarzucił zewnętrzny korektor, co proponował i co ja rekomenduję. '
                  'W żółtym polu na końcu wpisz swoją decyzję – wystarczy „tak", „nie" albo własne '
                  'brzmienie. Poprawki wprowadzę potem hurtem z tego dokumentu. '
                  'Rekomendacja to moja propozycja, nie rozstrzygnięcie: przy kilku pozycjach '
                  'korektor po prostu się mylił i piszę o tym wprost.')
    r.font.size = Pt(10)

    # ---- sekcja A
    doc.add_page_break()
    p = doc.add_paragraph()
    r = p.add_run('A. Decyzje zasadnicze')
    r.bold = True
    r.font.size = Pt(16)
    p = doc.add_paragraph()
    r = p.add_run('Cztery sprawy, które ważą więcej niż pojedyncze zdanie – '
                  'jedna z nich (Mk 1,15) rozstrzyga też pozycję 25 z listy B.')
    r.font.size = Pt(10)
    r.font.color.rgb = GREY

    for tytul, stan, dodatek, rekomendacja in ZASADNICZE:
        doc.add_paragraph().paragraph_format.space_after = Pt(2)
        t = nowa_tabela(doc)
        naglowek_pozycji(t, '', tytul)
        c = komorka(t, JEST_FILL)
        wiersz(c, 'STAN:', stan)
        if dodatek:
            c = komorka(t, WERSET_FILL)
            wiersz(c, '', dodatek, kursywa=True)
        c = komorka(t, REKO_FILL)
        wiersz(c, 'MOJA REKOMENDACJA:', rekomendacja)
        pole_decyzji(t, 4)

    # ---- sekcja B
    doc.add_page_break()
    p = doc.add_paragraph()
    r = p.add_run('B. Pozycje do decyzji (%d)' % len(pozycje))
    r.bold = True
    r.font.size = Pt(16)
    p = doc.add_paragraph()
    r = p.add_run('Kolejność jak w raporcie – alfabetycznie wg pliku studium.')
    r.font.size = Pt(10)
    r.font.color.rgb = GREY

    brak_reko = []
    for i, poz in enumerate(pozycje, 1):
        werdykt, uzas, nowe = znajdz_reko(poz['plik'], poz['gdzie'], poz.get('jest'))
        if not werdykt:
            brak_reko.append((poz['plik'], poz['gdzie'], (poz.get('jest') or '')[:45]))

        doc.add_paragraph().paragraph_format.space_after = Pt(2)
        t = nowa_tabela(doc)
        naglowek_pozycji(t, '%d.' % i, '%s – %s   (%s)' % (
            poz['plik'].replace('.json', ''), poz['gdzie'], poz.get('tytul', '')))

        c = komorka(t, JEST_FILL)
        wiersz(c, 'DZIŚ W APLIKACJI:', poz['tekst'])

        if poz['refs']:
            c = komorka(t, WERSET_FILL)
            for j, (ref, tekst) in enumerate(poz['refs']):
                wiersz(c, '%s (BE):' % (ref or '?'), tekst, kursywa=True, rozmiar=9,
                       pierwszy=(j == 0))

        c = komorka(t)
        wiersz(c, 'ZARZUT KOREKTORA:', poz.get('problem', '–'), rozmiar=9)
        wiersz(c, 'JEGO PROPOZYCJA:', poz.get('propozycja', '–'), rozmiar=9, pierwszy=False)

        c = komorka(t, REKO_FILL)
        wiersz(c, 'MOJA REKOMENDACJA:', werdykt or '(brak)')
        if uzas:
            wiersz(c, '', uzas, rozmiar=10, pierwszy=False)
        if nowe:
            wiersz(c, 'PROPONOWANE BRZMIENIE:', nowe, kursywa=True, pierwszy=False)

        pole_decyzji(t, 3)

    doc.save(OUT)
    print('Zapisano: %s' % OUT)
    print('Pozycji: %d  (sekcja A: %d)' % (len(pozycje), len(ZASADNICZE)))
    if brak_reko:
        print('BEZ REKOMENDACJI (%d) – do uzupelnienia w skrypcie:' % len(brak_reko))
        for k in brak_reko:
            print('   %s | %s | %s' % k)


if __name__ == '__main__':
    main()
